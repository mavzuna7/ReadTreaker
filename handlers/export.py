# handlers/export.py
from aiogram import Router
from aiogram.types import Message
from aiogram.filters import Command
from asgiref.sync import sync_to_async
import io
from aiogram.types import BufferedInputFile
from docx import Document
from docx.shared import Pt, RGBColor

router = Router()

@sync_to_async
def get_user_books_for_export(telegram_id):
    from books.models import TelegramUser, UserBook
    try:
        user = TelegramUser.objects.get(telegram_id=telegram_id)
        books = (
            UserBook.objects
            .filter(user=user)
            .select_related('book')
            .order_by('-status', '-date_read')
        )
        return list(books)
    except Exception as e:
        print("Ошибка экспорта:", e)
        return []

@router.message(Command("export"))
async def cmd_export(message: Message):
    books = await get_user_books_for_export(message.from_user.id)

    if not books:
        await message.answer(
            "📭 У вас пока нет книг для экспорта."
        )
        return
    
    read_books = [
    b for b in books
    if b.status == "read"
    ]

    want_books = [
        b for b in books
        if b.status == "want_to_read"
    ]

    ratings = [
        b.rating
        for b in read_books
        if b.rating
    ]

    avg_rating = (
        round(sum(ratings) / len(ratings), 1)
        if ratings else 0
    )

    doc = Document()

    p = doc.add_paragraph()
    run = p.add_run("📚 Моя библиотека ReadTracker")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = 1  

    doc.add_paragraph(
    f"📚 Всего книг: {len(books)}"
    )

    p = doc.add_paragraph()
    run = p.add_run("✅ Прочитано")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        f"⏳ Хочу прочитать: {len(want_books)}"
    )

    if avg_rating:
        doc.add_paragraph(
            f"⭐ Средняя оценка: {avg_rating}"
        )

    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run("✅ Прочитано")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    for num, ub in enumerate(read_books, start=1):

        p = doc.add_paragraph()
        run = p.add_run(f"{num}. {ub.book.title}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

        p = doc.add_paragraph()

        p.add_run("Автор: ").bold = True
        p.add_run(f"{ub.book.author}\n")

        if ub.book.genre:
            p.add_run("Жанр: ").bold = True
            p.add_run(f"{ub.book.genre}\n")

        if ub.book.year:
            p.add_run("Год: ").bold = True
            p.add_run(f"{ub.book.year}\n")

        if ub.rating:
            p.add_run("Оценка: ").bold = True
            p.add_run(f"{ub.rating}/5\n")

        if ub.date_start:
            p.add_run("Начато: ").bold = True
            p.add_run(f"{ub.date_start}\n")

        if ub.date_end:
            p.add_run("Завершено: ").bold = True
            p.add_run(f"{ub.date_end}\n")

        if ub.description:
            doc.add_paragraph(
                f"📖 Описание:\n{ub.description}"
            )

        if ub.review:
            doc.add_paragraph(
                f"💭 Мои впечатления:\n{ub.review}"
            )

        doc.add_paragraph(
            "─" * 40
        )

    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run("⏳ Хочу прочитать")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    for num, ub in enumerate(want_books, start=1):

        p = doc.add_paragraph()
        run = p.add_run(f"{num}. {ub.book.title}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

        p = doc.add_paragraph()

        p.add_run("Автор: ").bold = True
        p.add_run(f"{ub.book.author}\n")

        if ub.book.genre:
            p.add_run("Жанр: ").bold = True
            p.add_run(f"{ub.book.genre}\n")

        if ub.book.year:
            p.add_run("Год: ").bold = True
            p.add_run(f"{ub.book.year}\n")

        doc.add_paragraph(
            "─" * 40
        )

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    document = BufferedInputFile(
        file_stream.read(),
        filename="ReadTracker_Library.docx"
    )

    await message.answer_document(
        document=document,
        caption="📥 Экспорт библиотеки завершён!"
    )